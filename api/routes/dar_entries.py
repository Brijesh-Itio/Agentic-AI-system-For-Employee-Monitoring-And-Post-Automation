"""MODULE 7 extension (7.6-7.9) — Structured DAR entries, AI drafting, and
CSV/DOCX/PDF export + CSV import.

Entirely additive on top of module 7.1-7.4's narrative DAR generator
(ai/dar_generator.py) — nothing here changes generate_and_save_dar()'s
behaviour. dar_entries is its own table; the narrative DAR content is only
*read* (never written) here, to include alongside entries in exports.
"""
import csv
import io
import json
import logging
from datetime import date as date_type, datetime

from fastapi import APIRouter, Depends, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from api.auth import get_current_user
from api.database import DarEntry, DarReport, DarTemplate, Department, User, get_db
from api.schemas import DarEntryCreate, DarEntryDraftRequest, DarEntryOut, DarEntryUpdate

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/reports/dar", tags=["dar-entries"])


def _entry_out(row: DarEntry) -> DarEntryOut:
    return DarEntryOut(
        id=row.id,
        date=row.date,
        department_id=row.department_id,
        task=row.task,
        task_description=row.task_description,
        start_time=row.start_time,
        end_time=row.end_time,
        comment=row.comment,
        remarks=row.remarks,
        link=row.link,
        custom_fields=json.loads(row.custom_fields_json or "{}"),
        source=row.source,
        created_at=row.created_at,
        updated_at=row.updated_at,
    )


def _template_fields_for(db: Session, department_id: int | None) -> list[dict]:
    row = None
    if department_id is not None:
        row = db.query(DarTemplate).filter(DarTemplate.department_id == department_id).first()
    if row is None:
        row = db.query(DarTemplate).filter(DarTemplate.department_id.is_(None)).first()
    return json.loads(row.fields_json) if row else []


def _validate_custom_fields(db: Session, department_id: int | None, custom_fields: dict) -> None:
    fields = _template_fields_for(db, department_id)
    allowed_keys = {f["key"] for f in fields}
    unknown = set(custom_fields) - allowed_keys
    if unknown:
        raise HTTPException(
            status_code=422, detail=f"Unknown custom field(s) for this department: {sorted(unknown)}"
        )
    missing_required = [
        f["key"] for f in fields if f.get("required") and not str(custom_fields.get(f["key"], "")).strip()
    ]
    if missing_required:
        raise HTTPException(status_code=422, detail=f"Missing required field(s): {missing_required}")


# ── 7.6 Structured DAR entries (task log) — CRUD ──

@router.get("/date/{target_date}/entries", response_model=list[DarEntryOut])
def list_entries(
    target_date: date_type, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)
):
    rows = (
        db.query(DarEntry)
        .filter(DarEntry.user_id == current_user.id, DarEntry.date == target_date)
        .order_by(DarEntry.start_time.asc().nulls_last(), DarEntry.id.asc())
        .all()
    )
    return [_entry_out(r) for r in rows]


@router.post("/entries", response_model=DarEntryOut, status_code=201)
def create_entry(
    payload: DarEntryCreate, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)
):
    if payload.department_id is not None:
        if db.query(Department).filter(Department.id == payload.department_id).first() is None:
            raise HTTPException(status_code=404, detail=f"No department {payload.department_id}")
    _validate_custom_fields(db, payload.department_id, payload.custom_fields)

    row = DarEntry(
        user_id=current_user.id,
        date=payload.date,
        department_id=payload.department_id,
        task=payload.task,
        task_description=payload.task_description,
        start_time=payload.start_time,
        end_time=payload.end_time,
        comment=payload.comment,
        remarks=payload.remarks,
        link=payload.link,
        custom_fields_json=json.dumps(payload.custom_fields),
        source="manual",
        created_at=datetime.now(),
        updated_at=datetime.now(),
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    return _entry_out(row)


@router.patch("/entries/{entry_id}", response_model=DarEntryOut)
def update_entry(
    entry_id: int,
    payload: DarEntryUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    row = db.query(DarEntry).filter(DarEntry.id == entry_id).first()
    if row is None:
        raise HTTPException(status_code=404, detail=f"No entry {entry_id}")
    if row.user_id != current_user.id and current_user.role not in ("manager", "admin"):
        raise HTTPException(status_code=403, detail="Not authorised to edit this entry")

    updates = payload.model_dump(exclude_unset=True)
    if "custom_fields" in updates:
        _validate_custom_fields(db, row.department_id, updates["custom_fields"])
        row.custom_fields_json = json.dumps(updates.pop("custom_fields"))
    for field, value in updates.items():
        setattr(row, field, value)
    row.updated_at = datetime.now()
    db.commit()
    db.refresh(row)
    return _entry_out(row)


@router.delete("/entries/{entry_id}", status_code=204)
def delete_entry(
    entry_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)
):
    row = db.query(DarEntry).filter(DarEntry.id == entry_id).first()
    if row is None:
        raise HTTPException(status_code=404, detail=f"No entry {entry_id}")
    if row.user_id != current_user.id and current_user.role not in ("manager", "admin"):
        raise HTTPException(status_code=403, detail="Not authorised to delete this entry")
    db.delete(row)
    db.commit()


# ── 7.7 AI-assisted entry drafting ──

@router.post("/entries/draft", response_model=list[DarEntryOut])
def draft_entries(
    payload: DarEntryDraftRequest, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)
):
    """Asks Ollama to draft dar_entries rows matching the department's
    template, from the same day log 7.1 already builds. Never fabricates
    silently: an unparseable model response is a 502, not empty/fake rows."""
    department = db.query(Department).filter(Department.id == payload.department_id).first()
    if department is None:
        raise HTTPException(status_code=404, detail=f"No department {payload.department_id}")

    fields = _template_fields_for(db, payload.department_id)

    from ai.dar_generator import build_day_log
    from ai.ollama_client import generate

    day_log = build_day_log(payload.date, current_user.id)
    field_spec = ", ".join(f'"{f["key"]}" ({f["type"]})' for f in fields) or "(no custom fields defined)"
    prompt = (
        "You are drafting a structured task log for a Daily Activity Report, department: "
        f"{department.name}. Based on the raw activity log below, produce a JSON array of task "
        "entries. Each entry must be an object with exactly these base keys: "
        '"task" (short title), "task_description", "start_time" (HH:MM 24h), "end_time" (HH:MM 24h), '
        '"comment", "remarks", "link" (empty string if none), and a "custom_fields" object containing '
        f"these department-specific keys: {field_spec}.\n"
        "Base every entry strictly on the activity below — do not invent work that isn't reflected in "
        "it. Respond with ONLY the JSON array, no other text.\n\n"
        f"RAW ACTIVITY LOG:\n{day_log}"
    )

    raw = generate(prompt, fast=False)
    if raw is None:
        raise HTTPException(status_code=502, detail="Ollama unreachable or timed out while drafting entries")

    try:
        start = raw.index("[")
        end = raw.rindex("]") + 1
        drafted = json.loads(raw[start:end])
        if not isinstance(drafted, list):
            raise ValueError("model response was not a JSON array")
    except (ValueError, json.JSONDecodeError) as exc:
        logger.error("Could not parse drafted entries JSON: %s\nraw=%r", exc, raw)
        raise HTTPException(status_code=502, detail="Model response could not be parsed as JSON entries")

    created: list[DarEntry] = []
    for item in drafted:
        if not isinstance(item, dict) or not item.get("task"):
            continue
        custom_fields = item.get("custom_fields") or {}
        if not isinstance(custom_fields, dict):
            custom_fields = {}
        allowed_keys = {f["key"] for f in fields}
        custom_fields = {k: v for k, v in custom_fields.items() if k in allowed_keys}

        def _parse_clock(value):
            if not value:
                return None
            try:
                hh, mm = value.split(":")
                return datetime.combine(payload.date, datetime.min.time().replace(hour=int(hh), minute=int(mm)))
            except (ValueError, AttributeError):
                return None

        row = DarEntry(
            user_id=current_user.id,
            date=payload.date,
            department_id=payload.department_id,
            task=str(item.get("task"))[:500],
            task_description=item.get("task_description"),
            start_time=_parse_clock(item.get("start_time")),
            end_time=_parse_clock(item.get("end_time")),
            comment=item.get("comment"),
            remarks=item.get("remarks"),
            link=item.get("link") or None,
            custom_fields_json=json.dumps(custom_fields),
            source="ai_draft",
            created_at=datetime.now(),
            updated_at=datetime.now(),
        )
        db.add(row)
        created.append(row)

    if not created:
        raise HTTPException(status_code=502, detail="Model returned no usable entries")

    db.commit()
    for row in created:
        db.refresh(row)
    return [_entry_out(r) for r in created]


# ── 7.8 DAR export (CSV / DOCX / PDF) ──

def _department_name(db: Session, department_id: int | None) -> str:
    if department_id is None:
        return ""
    dept = db.query(Department).filter(Department.id == department_id).first()
    return dept.name if dept else ""


def _custom_field_keys(db: Session, entries: list[DarEntry]) -> list[str]:
    keys: list[str] = []
    seen = set()
    for e in entries:
        for k in json.loads(e.custom_fields_json or "{}"):
            if k not in seen:
                seen.add(k)
                keys.append(k)
    return keys


def _export_csv(db: Session, target_date: date_type, entries: list[DarEntry]) -> StreamingResponse:
    custom_keys = _custom_field_keys(db, entries)
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow(
        ["task", "task_description", "start_time", "end_time", "comment", "remarks", "link",
         "department", "source"] + custom_keys
    )
    for e in entries:
        custom = json.loads(e.custom_fields_json or "{}")
        writer.writerow(
            [
                e.task, e.task_description or "", e.start_time or "", e.end_time or "",
                e.comment or "", e.remarks or "", e.link or "",
                _department_name(db, e.department_id), e.source,
            ]
            + [custom.get(k, "") for k in custom_keys]
        )
    buffer.seek(0)
    filename = f"dar_{target_date.isoformat()}.csv"
    return StreamingResponse(
        iter([buffer.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _export_docx(db: Session, target_date: date_type, entries: list[DarEntry], narrative: str | None):
    from docx import Document

    doc = Document()
    doc.add_heading(f"Daily Activity Report — {target_date.isoformat()}", level=1)

    if narrative:
        doc.add_heading("Summary", level=2)
        for line in narrative.splitlines():
            doc.add_paragraph(line)

    doc.add_heading("Task Log", level=2)
    if not entries:
        doc.add_paragraph("No entries recorded for this date.")
    else:
        custom_keys = _custom_field_keys(db, entries)
        columns = ["Task", "Description", "Start", "End", "Comment", "Remarks", "Link",
                   "Department", "Source"] + custom_keys
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = "Light Grid Accent 1"
        for cell, header in zip(table.rows[0].cells, columns):
            cell.text = header
        for e in entries:
            custom = json.loads(e.custom_fields_json or "{}")
            values = [
                e.task, e.task_description or "", str(e.start_time or ""), str(e.end_time or ""),
                e.comment or "", e.remarks or "", e.link or "",
                _department_name(db, e.department_id), e.source,
            ] + [str(custom.get(k, "")) for k in custom_keys]
            cells = table.add_row().cells
            for cell, value in zip(cells, values):
                cell.text = value

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"dar_{target_date.isoformat()}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _export_pdf(db: Session, target_date: date_type, entries: list[DarEntry], narrative: str | None):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))
    styles = getSampleStyleSheet()
    story = [Paragraph(f"Daily Activity Report — {target_date.isoformat()}", styles["Title"]), Spacer(1, 12)]

    if narrative:
        story.append(Paragraph("Summary", styles["Heading2"]))
        for line in narrative.splitlines():
            if line.strip():
                story.append(Paragraph(line, styles["BodyText"]))
        story.append(Spacer(1, 12))

    story.append(Paragraph("Task Log", styles["Heading2"]))
    if not entries:
        story.append(Paragraph("No entries recorded for this date.", styles["BodyText"]))
    else:
        custom_keys = _custom_field_keys(db, entries)
        columns = ["Task", "Description", "Start", "End", "Comment", "Remarks", "Link",
                   "Department", "Source"] + custom_keys
        data = [columns]
        for e in entries:
            custom = json.loads(e.custom_fields_json or "{}")
            data.append(
                [
                    e.task, e.task_description or "", str(e.start_time or ""), str(e.end_time or ""),
                    e.comment or "", e.remarks or "", e.link or "",
                    _department_name(db, e.department_id), e.source,
                ]
                + [str(custom.get(k, "")) for k in custom_keys]
            )
        table = Table(data, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4f46e5")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, -1), 7),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(table)

    doc.build(story)
    buffer.seek(0)
    filename = f"dar_{target_date.isoformat()}.pdf"
    return StreamingResponse(
        buffer, media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/date/{target_date}/export")
def export_dar(
    target_date: date_type,
    format: str = "csv",
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if format not in ("csv", "docx", "pdf"):
        raise HTTPException(status_code=422, detail="format must be one of: csv, docx, pdf")

    entries = (
        db.query(DarEntry)
        .filter(DarEntry.user_id == current_user.id, DarEntry.date == target_date)
        .order_by(DarEntry.start_time.asc().nulls_last(), DarEntry.id.asc())
        .all()
    )
    report = (
        db.query(DarReport)
        .filter(DarReport.user_id == current_user.id, DarReport.date == target_date)
        .first()
    )
    narrative = report.content if report else None

    if format == "csv":
        return _export_csv(db, target_date, entries)
    if format == "docx":
        return _export_docx(db, target_date, entries, narrative)
    return _export_pdf(db, target_date, entries, narrative)


# ── 7.9 DAR import (CSV) ──

# "department" and "source" are metadata columns produced by 7.8's own CSV
# export — recognised specially so a round-trip export→import resolves the
# department by name instead of dumping it into custom_fields, and "source"
# is dropped rather than re-imported as a fake custom field (source is
# always "manual" for anything coming through this endpoint).
_BASE_COLUMNS = {"task", "task_description", "start_time", "end_time", "comment", "remarks", "link"}
_META_COLUMNS = {"department", "source"}


@router.post("/date/{target_date}/import", response_model=list[DarEntryOut])
async def import_dar_csv(
    target_date: date_type,
    file: UploadFile,
    department_id: int | None = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if department_id is not None and db.query(Department).filter(Department.id == department_id).first() is None:
        raise HTTPException(status_code=404, detail=f"No department {department_id}")

    raw = (await file.read()).decode("utf-8-sig")
    reader = csv.DictReader(io.StringIO(raw))
    if reader.fieldnames is None:
        raise HTTPException(status_code=422, detail="CSV has no header row")

    department_by_name = {d.name: d.id for d in db.query(Department).all()}

    def _parse_dt(value: str | None):
        if not value:
            return None
        try:
            return datetime.fromisoformat(value)
        except ValueError:
            return None

    created: list[DarEntry] = []
    for line_num, row in enumerate(reader, start=2):
        task = (row.get("task") or "").strip()
        if not task:
            continue  # 7.9: never lose data silently, but a row with no task has nothing to import
        row_department_id = department_by_name.get((row.get("department") or "").strip(), department_id)
        custom_fields = {
            k: v for k, v in row.items()
            if k not in _BASE_COLUMNS and k not in _META_COLUMNS and v not in (None, "")
        }
        entry = DarEntry(
            user_id=current_user.id,
            date=target_date,
            department_id=row_department_id,
            task=task,
            task_description=row.get("task_description") or None,
            start_time=_parse_dt(row.get("start_time")),
            end_time=_parse_dt(row.get("end_time")),
            comment=row.get("comment") or None,
            remarks=row.get("remarks") or None,
            link=row.get("link") or None,
            custom_fields_json=json.dumps(custom_fields),
            source="manual",
            created_at=datetime.now(),
            updated_at=datetime.now(),
        )
        db.add(entry)
        created.append(entry)

    if not created:
        raise HTTPException(status_code=422, detail="No importable rows found (every row was missing 'task')")

    db.commit()
    for row in created:
        db.refresh(row)
    return [_entry_out(r) for r in created]
